# utils.py

import re
from io import BytesIO
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from openai import OpenAI


def chunk_embed_store_transcript(transcript_text, persist_dir="./chroma_db"):
    """
    Splits transcript into chunks, generates embeddings,
    and stores them in a persistent Chroma vector database.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", ""],
        chunk_size=1000,
        chunk_overlap=100,
    )
    chunks = text_splitter.split_text(transcript_text)

    embeddings_model = OpenAIEmbeddings(model="text-embedding-3-small")
    vectordb = Chroma.from_texts(
        texts=chunks,
        embedding=embeddings_model,
        persist_directory=persist_dir,
        collection_name="decode_transcripts",
    )
    vectordb.persist()
    return vectordb


def build_retriever(persist_dir="./chroma_db"):
    """
    Loads the Chroma vector database and returns a retriever.
    Prefer MMR to diversify results; otherwise fall back to k search.
    """
    embeddings_model = OpenAIEmbeddings(model="text-embedding-3-small")
    vectordb = Chroma(
        persist_directory=persist_dir,
        embedding_function=embeddings_model,
        collection_name="decode_transcripts",
    )
    try:
        retriever = vectordb.as_retriever(
            search_type="mmr",
            search_kwargs={"k": 10, "fetch_k": 25, "lambda_mult": 0.5},
        )
    except Exception:
        retriever = vectordb.as_retriever(search_kwargs={"k": 8})
    return retriever


def get_llm_client():
    """Returns an OpenAI client instance."""
    return OpenAI()


def generate_insights(client, question, docs):
    """
    Generates concise, numbered insights grounded in the provided excerpts.
    - No headings/section titles.
    - Each insight = short headline + concise explanation.
    - If using quotes, they must be verbatim from the excerpts.
    - Do not repeat the headline inside the explanation.
    """
    context = "\n\n".join(doc.page_content for doc in docs)

    system_msg = (
        "You are an expert qualitative research analyst. "
        "Your job is to produce clear, actionable, numbered insights grounded only in the provided excerpts."
    )

    user_msg = f"""
Research question:
{question}

Excerpts (verbatim):
{context}

Write 3–6 insights, numbered "1., 2., 3., ...", with this exact format per item:

- Start with a short headline (max 12 words), then an em dash (—), then a single concise explanation (2–3 sentences).
- Do NOT include section headers or titles (e.g., '###', 'Insights', 'Summary').
- Do NOT repeat the headline inside the explanation.
- Only use quotation marks when copying verbatim from the excerpts. Never invent quotes.
- Prefer patterns, tensions, contradictions, and actionable implications.
- Avoid vague language; be concrete and specific.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.1,
        max_tokens=700,
    )
    return response.choices[0].message.content


def split_insights_into_points(insight_text):
    """
    Splits the AI-generated insights text into numbered points.
    Handles '1.' and '1)' markers and filters out heading-like parts.
    """
    pattern = r'(?:^|\n)\s*(?:\d+\.\s|\d+\)\s)'
    parts = re.split(pattern, insight_text)
    points = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if p.startswith("#") or p.lower().startswith(("insight", "summary")):
            continue
        points.append(p)
    return points


def extract_insight_summaries(insights):
    """
    Extract concise summaries from insight points.
    Use first sentence or first ~10 words as a fallback.
    """
    summaries = []
    for insight in insights:
        first_sentence = insight.split('.', 1)[0].strip()
        first_words = ' '.join(insight.split()[:10])
        summary = first_sentence if len(first_sentence) > 5 else first_words
        summaries.append(summary)
    return summaries


def find_supporting_quotes(points, all_chunks, embeddings_model, top_k=2):
    """
    Hybrid approach for alignment with inline quotes:
      1) Extract quoted phrases from the insight and exact-match them in chunks.
      2) If not enough matches, fall back to semantic similarity.
    Returns a list of lists of chunk strings.
    """
    # Precompute chunk embeddings once
    chunk_embeds = embeddings_model.embed_documents(all_chunks)
    results = []

    for point in points:
        # Extract inline quotes (curly or straight). Ignore very short snippets.
        quoted_bits = []
        for m in re.findall(r'“([^”]+)”|"([^"]+)"', point):
            frag = (m[0] or m[1]).strip()
            if len(frag) >= 8:
                quoted_bits.append(frag)

        matched = []

        # 1) Exact (case-insensitive) search of quoted phrases
        if quoted_bits:
            for qb in quoted_bits:
                qb_norm = re.sub(r'\s+', ' ', qb).strip().lower()
                best_idx = None
                for i, ch in enumerate(all_chunks):
                    ch_norm = re.sub(r'\s+', ' ', ch).strip().lower()
                    if qb_norm in ch_norm:
                        best_idx = i
                        break
                if best_idx is not None:
                    cand = all_chunks[best_idx]
                    if cand not in matched:
                        matched.append(cand)
                        if len(matched) >= top_k:
                            break

        # 2) Fall back to semantic similarity
        if len(matched) < top_k:
            query_text = " ".join(quoted_bits) if quoted_bits else point
            q_embed = embeddings_model.embed_query(query_text)
            sims = cosine_similarity([q_embed], chunk_embeds)[0]
            for idx in sims.argsort()[::-1]:
                cand = all_chunks[idx]
                if cand not in matched:
                    matched.append(cand)
                    if len(matched) >= top_k:
                        break

        results.append(matched[:top_k])

    return results


def export_to_word(chat_history, supporting_quotes):
    """
    Exports the chat history and supporting quotes to a Word document (in-memory).
    """
    doc = Document()
    doc.add_heading("Decode Findings Export", 0)
    for i, (q, a) in enumerate(chat_history):
        doc.add_heading(f"Q{i+1}: {q}", level=1)
        # Keep original points in export; UI handles repetition removal
        points = split_insights_into_points(a)
        for j, point in enumerate(points):
            doc.add_heading(f"Insight {j+1}:", level=2)
            doc.add_paragraph(point)
            doc.add_heading("Supporting Quotes:", level=3)
            for quote in supporting_quotes[i][j]:
                doc.add_paragraph(f"- {quote}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
